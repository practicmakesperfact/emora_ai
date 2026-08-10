"""
Emora Backend - Document Service
Handles file saving, text extraction from PDF/DOCX/TXT/MD files,
and persisting document metadata to the database.
"""

import os
import uuid
from pathlib import Path
from typing import Sequence

from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.exceptions import NotFoundException, AppException
from app.core.logging import get_logger
from app.models.document import KnowledgeDocument
from app.repositories.document import DocumentRepository

logger = get_logger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class DocumentService:
    """Service for uploading, extracting, and managing knowledge documents."""

    def __init__(self, db: AsyncSession) -> None:
        self._db = db
        self._repo = DocumentRepository(db)
        self._upload_dir = Path(settings.UPLOAD_DIR)
        self._upload_dir.mkdir(parents=True, exist_ok=True)

    async def upload_document(
        self,
        file: UploadFile,
        title: str,
        author: str | None = None,
        source: str | None = None,
    ) -> KnowledgeDocument:
        """
        Save the uploaded file, extract its text content, and persist metadata.
        """
        suffix = Path(file.filename or "file").suffix.lower()
        if suffix not in ALLOWED_EXTENSIONS:
            raise AppException(
                status_code=400,
                message=f"Unsupported file type '{suffix}'. Allowed: PDF, DOCX, TXT, MD.",
            )

        # Save file with unique name
        unique_name = f"{uuid.uuid4().hex}{suffix}"
        file_path = self._upload_dir / unique_name

        content_bytes = await file.read()
        file_path.write_bytes(content_bytes)

        # Extract text content
        text_content = self._extract_text(file_path, suffix, content_bytes)

        document = KnowledgeDocument(
            title=title,
            author=author,
            source=source,
            file_name=file.filename or unique_name,
            file_path=str(file_path),
            content=text_content,
        )

        created_doc = await self._repo.create_document(document)
        logger.info(
            "Document uploaded and saved",
            doc_id=created_doc.id,
            title=title,
            file=file.filename,
        )
        return created_doc

    async def list_documents(self, skip: int = 0, limit: int = 100) -> Sequence[KnowledgeDocument]:
        """List all knowledge documents."""
        return await self._repo.get_all(skip=skip, limit=limit)

    async def delete_document(self, document_id: int) -> None:
        """Delete a document from the database and remove the file from disk."""
        document = await self._repo.get_by_id(document_id)
        if not document:
            raise NotFoundException(f"Document {document_id} not found.")

        # Remove file from disk
        file_path = Path(document.file_path)
        if file_path.exists():
            file_path.unlink()

        await self._repo.delete_document(document)
        logger.info("Document deleted", doc_id=document_id)

    def _extract_text(self, file_path: Path, suffix: str, content_bytes: bytes) -> str:
        """Extract plain text from PDF, DOCX, TXT, or MD files."""
        try:
            if suffix == ".pdf":
                from pypdf import PdfReader
                from io import BytesIO
                reader = PdfReader(BytesIO(content_bytes))
                return "\n".join(
                    page.extract_text() or "" for page in reader.pages
                )
            elif suffix == ".docx":
                from docx import Document
                from io import BytesIO
                doc = Document(BytesIO(content_bytes))
                return "\n".join(para.text for para in doc.paragraphs if para.text)
            elif suffix in {".txt", ".md"}:
                return content_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.error("Text extraction failed", error=str(e), file=str(file_path))
        return ""
